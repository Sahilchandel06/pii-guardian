import csv
import io
import json
import mimetypes
import os
from collections import defaultdict
from typing import Any

from app.services.pii_detector import detect_pii
from app.services.sanitizer import sanitize_text

SUPPORTED_EXTENSIONS = {
    ".sql",
    ".csv",
    ".json",
    ".pdf",
    ".docx",
    ".txt",
    ".png",
    ".jpg",
    ".jpeg",
    ".xlsx",
    ".xlsm",
    ".xltx",
    ".xltm",
    ".xls",
}

EXCEL_OPENXML_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
GENERIC_TEXT_MIME_PREFIXES = ("text/",)
GENERIC_STRUCTURED_MIME_TYPES = {
    "application/json",
    "application/xml",
    "application/x-yaml",
    "application/yaml",
}
_paddle_ocr = None


def _looks_like_text(file_bytes: bytes) -> bool:
    sample = file_bytes[:4096]
    if not sample:
        return True
    try:
        sample.decode("utf-8")
        return True
    except UnicodeDecodeError:
        pass
    allowed_controls = {9, 10, 13}
    printable = sum((32 <= b <= 126) or (b in allowed_controls) for b in sample)
    return (printable / len(sample)) >= 0.9


def is_supported_file(filename: str, content_type: str | None = None, file_bytes: bytes | None = None) -> bool:
    ext = os.path.splitext(filename.lower())[1]
    if ext in SUPPORTED_EXTENSIONS:
        return True
    normalized_type = (content_type or "").split(";")[0].strip().lower()
    if normalized_type.startswith(GENERIC_TEXT_MIME_PREFIXES):
        return True
    if normalized_type in GENERIC_STRUCTURED_MIME_TYPES:
        return True
    if file_bytes is not None and _looks_like_text(file_bytes):
        return True
    return False


def _collect_counts(chunks: list[list[dict[str, Any]]]) -> dict[str, int]:
    counts: dict[str, int] = defaultdict(int)
    for findings in chunks:
        for finding in findings:
            counts[finding["entity_type"]] += 1
    return dict(counts)


def _sanitize_chunk(text: str, mode: str) -> tuple[str, list[dict[str, Any]]]:
    findings = detect_pii(text)
    sanitized, _ = sanitize_text(text, findings, mode=mode)
    return sanitized, findings


def _ocr_lang() -> str:
    # PaddleOCR language code, e.g. en, ch, hi.
    return os.getenv("OCR_LANG", "en")


def _tesseract_lang() -> str:
    lang = _ocr_lang().lower()
    mapping = {"en": "eng"}
    return mapping.get(lang, lang)


def _configure_paddle_cache() -> None:
    cache_home = os.getenv("PADDLE_PDX_CACHE_HOME")
    if not cache_home:
        cache_home = os.path.abspath(os.path.join(os.getcwd(), ".paddlex"))
        os.environ["PADDLE_PDX_CACHE_HOME"] = cache_home
    os.makedirs(cache_home, exist_ok=True)
    # Avoid home-dir permission issues from matplotlib used by downstream OCR deps.
    mpl_config_dir = os.getenv("MPLCONFIGDIR")
    if not mpl_config_dir:
        mpl_config_dir = os.path.join(cache_home, "mplconfig")
        os.environ["MPLCONFIGDIR"] = mpl_config_dir
    os.makedirs(mpl_config_dir, exist_ok=True)


def _get_paddle_ocr():
    global _paddle_ocr
    if _paddle_ocr is not None:
        return _paddle_ocr

    _configure_paddle_cache()
    from paddleocr import PaddleOCR

    use_angle_cls = os.getenv("OCR_USE_ANGLE_CLS", "true").lower() == "true"
    _paddle_ocr = PaddleOCR(
        use_angle_cls=use_angle_cls,
        lang=_ocr_lang(),
        show_log=False,
    )
    return _paddle_ocr


def _extract_ocr_segments_tesseract(image) -> tuple[list[dict[str, Any]], str]:
    import pytesseract

    ocr = pytesseract.image_to_data(
        image,
        output_type=pytesseract.Output.DICT,
        lang=_tesseract_lang(),
        config="--oem 3 --psm 6",
    )

    segments: list[dict[str, Any]] = []
    cursor = 0
    parts: list[str] = []

    for idx, raw_text in enumerate(ocr.get("text", [])):
        text = (raw_text or "").strip()
        if not text:
            continue
        try:
            confidence = float(ocr.get("conf", ["-1"])[idx])
        except (TypeError, ValueError):
            confidence = -1
        if confidence < 0:
            continue

        if parts:
            parts.append(" ")
            cursor += 1
        start = cursor
        parts.append(text)
        cursor += len(text)

        segments.append(
            {
                "start": start,
                "end": cursor,
                "left": int(ocr["left"][idx]),
                "top": int(ocr["top"][idx]),
                "width": int(ocr["width"][idx]),
                "height": int(ocr["height"][idx]),
            }
        )

    return segments, "".join(parts)


def _extract_ocr_segments(image) -> tuple[list[dict[str, Any]], str]:
    try:
        import numpy as np
    except Exception as exc:
        raise RuntimeError("NumPy is required for PaddleOCR image processing.") from exc

    try:
        ocr_engine = _get_paddle_ocr()
    except Exception as exc:
        try:
            return _extract_ocr_segments_tesseract(image)
        except Exception:
            raise RuntimeError(
                "PaddleOCR is unavailable. Install paddleocr and paddlepaddle, then retry."
            ) from exc
    try:
        result = ocr_engine.ocr(np.array(image.convert("RGB")), cls=True)
    except Exception:
        return _extract_ocr_segments_tesseract(image)
    lines = result[0] if isinstance(result, list) and result else []

    segments: list[dict[str, Any]] = []
    cursor = 0
    parts: list[str] = []

    for line in lines or []:
        if not isinstance(line, list) or len(line) < 2:
            continue

        box = line[0]
        text_block = line[1]
        if not text_block or len(text_block) < 1:
            continue

        text = str(text_block[0] or "").strip()
        if not text:
            continue
        try:
            confidence = float(text_block[1]) if len(text_block) > 1 else 0.0
        except (TypeError, ValueError):
            confidence = 0.0
        if confidence < 0:
            continue

        try:
            x_values = [int(point[0]) for point in box]
            y_values = [int(point[1]) for point in box]
        except Exception:
            continue
        left = min(x_values)
        right = max(x_values)
        top = min(y_values)
        bottom = max(y_values)

        if parts:
            parts.append(" ")
            cursor += 1
        start = cursor
        parts.append(text)
        cursor += len(text)

        segments.append(
            {
                "start": start,
                "end": cursor,
                "left": left,
                "top": top,
                "width": max(1, right - left),
                "height": max(1, bottom - top),
            }
        )

    return segments, "".join(parts)


def _image_to_string_with_fallback(image) -> str:
    _, plain_text = _extract_ocr_segments(image)
    return plain_text


def _qr_scan_enabled() -> bool:
    return os.getenv("USE_QR_SCAN", "true").lower() == "true"


def _decode_qr_payloads(image) -> list[dict[str, Any]]:
    if not _qr_scan_enabled():
        return []
    try:
        from pyzbar.pyzbar import decode
    except Exception:
        return []

    payloads: list[dict[str, Any]] = []
    try:
        decoded_items = decode(image)
    except Exception:
        return []

    for item in decoded_items:
        raw_data = getattr(item, "data", b"")
        if isinstance(raw_data, bytes):
            text = raw_data.decode("utf-8", errors="ignore").strip()
        else:
            text = str(raw_data).strip()
        if not text:
            continue

        left = 0
        top = 0
        width = 0
        height = 0

        rect = getattr(item, "rect", None)
        if rect is not None:
            left = int(getattr(rect, "left", 0))
            top = int(getattr(rect, "top", 0))
            width = int(getattr(rect, "width", 0))
            height = int(getattr(rect, "height", 0))

        if width <= 0 or height <= 0:
            polygon = getattr(item, "polygon", None) or []
            try:
                xs = [int(point.x) for point in polygon]
                ys = [int(point.y) for point in polygon]
                if xs and ys:
                    left = min(xs)
                    top = min(ys)
                    width = max(xs) - left
                    height = max(ys) - top
            except Exception:
                pass

        payloads.append(
            {
                "text": text,
                "left": max(0, left),
                "top": max(0, top),
                "width": max(1, width),
                "height": max(1, height),
            }
        )

    return payloads


def _analyze_qr_payloads(payloads: list[dict[str, Any]], mode: str) -> tuple[list[dict[str, Any]], list[str], list[dict[str, Any]]]:
    findings: list[dict[str, Any]] = []
    sanitized_blocks: list[str] = []
    pii_boxes: list[dict[str, Any]] = []

    for idx, payload in enumerate(payloads, start=1):
        qr_text = str(payload.get("text", "")).strip()
        if not qr_text:
            continue
        qr_findings = detect_pii(qr_text)
        if not qr_findings:
            continue
        pii_boxes.append(payload)
        sanitized_qr, _ = sanitize_text(qr_text, qr_findings, mode=mode)
        sanitized_blocks.append(f"[QR_{idx}] {sanitized_qr}")

        for finding in qr_findings:
            merged = dict(finding)
            merged["source"] = "qr"
            findings.append(merged)

    return findings, sanitized_blocks, pii_boxes


def _sanitize_image(image, mode: str):
    from PIL import ImageDraw

    segments, plain_text = _extract_ocr_segments(image)
    sanitized_ocr_text, ocr_findings = _sanitize_chunk(plain_text, mode)
    qr_payloads = _decode_qr_payloads(image)
    qr_findings, qr_sanitized_blocks, qr_boxes = _analyze_qr_payloads(qr_payloads, mode)

    findings = ocr_findings + qr_findings
    sanitized_text = sanitized_ocr_text
    if qr_sanitized_blocks:
        sanitized_text = f"{sanitized_text}\n" if sanitized_text else ""
        sanitized_text += "\n".join(qr_sanitized_blocks)

    redacted = image.copy().convert("RGB")
    draw = ImageDraw.Draw(redacted)
    for segment in segments:
        overlaps = any(
            finding["start"] < segment["end"] and segment["start"] < finding["end"]
            for finding in ocr_findings
        )
        if overlaps:
            x0 = segment["left"]
            y0 = segment["top"]
            x1 = x0 + max(1, segment["width"])
            y1 = y0 + max(1, segment["height"])
            draw.rectangle([x0, y0, x1, y1], fill="black")

    for box in qr_boxes:
        x0 = int(box["left"])
        y0 = int(box["top"])
        x1 = x0 + max(1, int(box["width"]))
        y1 = y0 + max(1, int(box["height"]))
        draw.rectangle([x0, y0, x1, y1], fill="black")

    # attempt to detect faces; if found, blackout face regions only
    try:
        import cv2
        import numpy as np

        # convert PIL -> OpenCV BGR
        cv_img = cv2.cvtColor(np.array(redacted), cv2.COLOR_RGB2BGR)
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
        faces = face_cascade.detectMultiScale(cv_img, scaleFactor=1.1, minNeighbors=5)
        # filter face boxes by reasonable aspect ratio and location
        valid_faces = []
        for (x, y, w, h) in faces:
            if w == 0 or h == 0:
                continue
            ratio = w / h
            # face roughly square and located in upper half of image
            if 0.5 < ratio < 1.5 and y < redacted.height * 0.6:
                valid_faces.append((x, y, w, h))
        if valid_faces:
            # choose largest box if multiple
            x, y, w, h = max(valid_faces, key=lambda b: b[2] * b[3])
            draw.rectangle([x, y, x + w, y + h], fill="black")
    except Exception:
        # Keep OCR-based text redaction even when face detection is unavailable.
        pass

    raw_text = plain_text
    if qr_payloads:
        qr_text_lines = [f"[QR_{idx}] {item['text']}" for idx, item in enumerate(qr_payloads, start=1)]
        raw_text = f"{raw_text}\n" if raw_text else ""
        raw_text += "\n".join(qr_text_lines)

    return redacted, raw_text, sanitized_text, findings


def _extract_from_excel_openxml(file_bytes: bytes) -> tuple[str, str]:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), data_only=False)
    chunks: list[str] = []
    for sheet in wb.worksheets:
        chunks.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                chunks.append(" | ".join(values))
    return "\n".join(chunks), "excel"


def _extract_from_xls(file_bytes: bytes) -> tuple[str, str]:
    import xlrd

    wb = xlrd.open_workbook(file_contents=file_bytes)
    chunks: list[str] = []
    for sheet in wb.sheets():
        chunks.append(f"[Sheet: {sheet.name}]")
        for row_idx in range(sheet.nrows):
            row_values = [str(sheet.cell_value(row_idx, col_idx) or "") for col_idx in range(sheet.ncols)]
            if any(row_values):
                chunks.append(" | ".join(row_values))
    return "\n".join(chunks), "excel_xls"


def extract_text_from_file(filename: str, file_bytes: bytes) -> tuple[str, str]:
    ext = os.path.splitext(filename.lower())[1]

    if ext in {".txt", ".sql"}:
        return file_bytes.decode("utf-8", errors="ignore"), ext[1:]

    if ext == ".json":
        decoded = file_bytes.decode("utf-8", errors="ignore")
        parsed = json.loads(decoded)
        return json.dumps(parsed, indent=2, ensure_ascii=False), "json"

    if ext == ".csv":
        decoded = file_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(decoded))
        lines = [",".join(row) for row in reader]
        return "\n".join(lines), "csv"

    if ext == ".pdf":
        import fitz
        from PIL import Image

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        chunks: list[str] = []
        for page in doc:
            text = (page.get_text("text") or "").strip()
            if text:
                chunks.append(text)
                continue

            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            chunks.append(_image_to_string_with_fallback(image))
        return "\n\n".join(chunks), "pdf_ocr"

    if ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        lines: list[str] = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines), "docx"

    if ext in IMAGE_EXTENSIONS:
        from PIL import Image

        image = Image.open(io.BytesIO(file_bytes))
        _, plain_text, _, _ = _sanitize_image(image, mode="mask")
        return plain_text, "ocr_image"

    if ext in EXCEL_OPENXML_EXTENSIONS:
        return _extract_from_excel_openxml(file_bytes)

    if ext == ".xls":
        return _extract_from_xls(file_bytes)

    if _looks_like_text(file_bytes):
        return file_bytes.decode("utf-8", errors="ignore"), "generic_text"

    raise ValueError(f"Unsupported file type: {ext or '[no extension]'}")


def _sanitize_json(obj: Any, mode: str, findings_bucket: list[list[dict[str, Any]]]) -> Any:
    if isinstance(obj, str):
        sanitized, findings = _sanitize_chunk(obj, mode)
        findings_bucket.append(findings)
        return sanitized
    if isinstance(obj, list):
        return [_sanitize_json(item, mode, findings_bucket) for item in obj]
    if isinstance(obj, dict):
        return {key: _sanitize_json(value, mode, findings_bucket) for key, value in obj.items()}
    return obj


def sanitize_file_preserving_format(
    filename: str,
    file_bytes: bytes,
    mode: str,
) -> dict[str, Any]:
    ext = os.path.splitext(filename.lower())[1]
    media_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"

    if ext in {".txt", ".sql"}:
        raw_text = file_bytes.decode("utf-8", errors="ignore")
        sanitized_text, findings = _sanitize_chunk(raw_text, mode)
        return {
            "sanitized_text": sanitized_text,
            "sanitized_original_bytes": sanitized_text.encode("utf-8"),
            "sanitized_original_ext": ext,
            "sanitized_original_media_type": "text/plain",
            "findings": findings,
            "summary": _collect_counts([findings]),
            "parser": ext[1:],
        }

    if ext == ".json":
        decoded = file_bytes.decode("utf-8", errors="ignore")
        payload = json.loads(decoded)
        findings_bucket: list[list[dict[str, Any]]] = []
        sanitized_payload = _sanitize_json(payload, mode, findings_bucket)
        sanitized_json = json.dumps(sanitized_payload, indent=2, ensure_ascii=False)
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": sanitized_json,
            "sanitized_original_bytes": sanitized_json.encode("utf-8"),
            "sanitized_original_ext": ".json",
            "sanitized_original_media_type": "application/json",
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "json",
        }

    if ext == ".csv":
        decoded = file_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(decoded))
        output = io.StringIO()
        writer = csv.writer(output, lineterminator="\n")
        findings_bucket: list[list[dict[str, Any]]] = []

        for row in reader:
            clean_row: list[str] = []
            for cell in row:
                sanitized_cell, findings = _sanitize_chunk(cell, mode)
                findings_bucket.append(findings)
                clean_row.append(sanitized_cell)
            writer.writerow(clean_row)

        sanitized_csv = output.getvalue()
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": sanitized_csv,
            "sanitized_original_bytes": sanitized_csv.encode("utf-8"),
            "sanitized_original_ext": ".csv",
            "sanitized_original_media_type": "text/csv",
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "csv",
        }

    if ext in EXCEL_OPENXML_EXTENSIONS:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_bytes), data_only=False)
        findings_bucket: list[list[dict[str, Any]]] = []
        preview_lines: list[str] = []

        for sheet in wb.worksheets:
            preview_lines.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows():
                row_preview: list[str] = []
                for cell in row:
                    value = cell.value
                    if isinstance(value, str):
                        sanitized_cell, findings = _sanitize_chunk(value, mode)
                        findings_bucket.append(findings)
                        cell.value = sanitized_cell
                        row_preview.append(sanitized_cell)
                    else:
                        row_preview.append("" if value is None else str(value))
                if any(row_preview):
                    preview_lines.append(" | ".join(row_preview))

        output = io.BytesIO()
        wb.save(output)
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": "\n".join(preview_lines),
            "sanitized_original_bytes": output.getvalue(),
            "sanitized_original_ext": ext,
            "sanitized_original_media_type": media_type,
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "excel",
        }

    if ext == ".xls":
        import xlrd
        import xlwt

        workbook = xlrd.open_workbook(file_contents=file_bytes)
        out_workbook = xlwt.Workbook()
        findings_bucket: list[list[dict[str, Any]]] = []
        preview_lines: list[str] = []

        for sheet_idx in range(workbook.nsheets):
            in_sheet = workbook.sheet_by_index(sheet_idx)
            out_sheet = out_workbook.add_sheet(in_sheet.name[:31] or f"Sheet{sheet_idx + 1}")
            preview_lines.append(f"[Sheet: {in_sheet.name}]")

            for row_idx in range(in_sheet.nrows):
                row_preview: list[str] = []
                for col_idx in range(in_sheet.ncols):
                    value = in_sheet.cell_value(row_idx, col_idx)
                    if isinstance(value, str):
                        sanitized_cell, findings = _sanitize_chunk(value, mode)
                        findings_bucket.append(findings)
                        out_sheet.write(row_idx, col_idx, sanitized_cell)
                        row_preview.append(sanitized_cell)
                    else:
                        out_sheet.write(row_idx, col_idx, value)
                        row_preview.append("" if value is None else str(value))
                if any(row_preview):
                    preview_lines.append(" | ".join(row_preview))

        output = io.BytesIO()
        out_workbook.save(output)
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": "\n".join(preview_lines),
            "sanitized_original_bytes": output.getvalue(),
            "sanitized_original_ext": ".xls",
            "sanitized_original_media_type": "application/vnd.ms-excel",
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "excel_xls",
        }

    if ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        findings_bucket: list[list[dict[str, Any]]] = []

        for paragraph in doc.paragraphs:
            if paragraph.text:
                sanitized_paragraph, findings = _sanitize_chunk(paragraph.text, mode)
                paragraph.text = sanitized_paragraph
                findings_bucket.append(findings)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        sanitized_cell, findings = _sanitize_chunk(cell.text, mode)
                        cell.text = sanitized_cell
                        findings_bucket.append(findings)

        output = io.BytesIO()
        doc.save(output)
        preview_text, _ = extract_text_from_file(filename, output.getvalue())
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": preview_text,
            "sanitized_original_bytes": output.getvalue(),
            "sanitized_original_ext": ".docx",
            "sanitized_original_media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "docx",
        }

    if ext in IMAGE_EXTENSIONS:
        from PIL import Image

        image = Image.open(io.BytesIO(file_bytes))
        redacted, _, sanitized_text, findings = _sanitize_image(image, mode)

        output = io.BytesIO()
        image_format = "PNG" if ext == ".png" else "JPEG"
        redacted.save(output, format=image_format)
        return {
            "sanitized_text": sanitized_text,
            "sanitized_original_bytes": output.getvalue(),
            "sanitized_original_ext": ext,
            "sanitized_original_media_type": media_type,
            "findings": findings,
            "summary": _collect_counts([findings]),
            "parser": "ocr_image",
        }

    if ext == ".pdf":
        import fitz
        from PIL import Image

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        findings_bucket: list[list[dict[str, Any]]] = []
        sanitized_pages: list[str] = []

        for page in doc:
            page_text = (page.get_text("text") or "").strip()
            qr_payloads: list[dict[str, Any]] = []
            qr_pix = None

            if _qr_scan_enabled():
                try:
                    qr_pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    qr_image = Image.frombytes("RGB", (qr_pix.width, qr_pix.height), qr_pix.samples)
                    qr_payloads = _decode_qr_payloads(qr_image)
                except Exception:
                    qr_payloads = []
                    qr_pix = None

            if page_text:
                text_findings = detect_pii(page_text)
                sanitized_text, _ = sanitize_text(page_text, text_findings, mode=mode)
                qr_findings, qr_sanitized_blocks, qr_boxes = _analyze_qr_payloads(qr_payloads, mode)
                findings = text_findings + qr_findings
                if qr_sanitized_blocks:
                    sanitized_text = f"{sanitized_text}\n" if sanitized_text else ""
                    sanitized_text += "\n".join(qr_sanitized_blocks)
                sanitized_pages.append(sanitized_text)
                findings_bucket.append(findings)

                unique_values = []
                seen_values = set()
                for finding in text_findings:
                    value = finding.get("value", "")
                    if value and value not in seen_values:
                        seen_values.add(value)
                        unique_values.append(value)

                for value in unique_values:
                    for rect in page.search_for(value):
                        page.add_redact_annot(rect, fill=(0, 0, 0))

                if qr_boxes and qr_pix is not None:
                    x_scale = page.rect.width / max(1, qr_pix.width)
                    y_scale = page.rect.height / max(1, qr_pix.height)
                    for box in qr_boxes:
                        x0 = box["left"] * x_scale
                        y0 = box["top"] * y_scale
                        x1 = (box["left"] + max(1, box["width"])) * x_scale
                        y1 = (box["top"] + max(1, box["height"])) * y_scale
                        page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(0, 0, 0))
                page.apply_redactions()
                continue

            # OCR fallback for scanned/image-only pages.
            pix = qr_pix if qr_pix is not None else page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            segments, page_ocr_text = _extract_ocr_segments(image)
            ocr_findings = detect_pii(page_ocr_text)
            sanitized_text, _ = sanitize_text(page_ocr_text, ocr_findings, mode=mode)
            qr_findings, qr_sanitized_blocks, qr_boxes = _analyze_qr_payloads(qr_payloads, mode)
            findings = ocr_findings + qr_findings
            if qr_sanitized_blocks:
                sanitized_text = f"{sanitized_text}\n" if sanitized_text else ""
                sanitized_text += "\n".join(qr_sanitized_blocks)
            sanitized_pages.append(sanitized_text)
            findings_bucket.append(findings)

            x_scale = page.rect.width / max(1, pix.width)
            y_scale = page.rect.height / max(1, pix.height)

            for segment in segments:
                overlaps = any(
                    finding["start"] < segment["end"] and segment["start"] < finding["end"]
                    for finding in ocr_findings
                )
                if not overlaps:
                    continue
                x0 = segment["left"] * x_scale
                y0 = segment["top"] * y_scale
                x1 = (segment["left"] + max(1, segment["width"])) * x_scale
                y1 = (segment["top"] + max(1, segment["height"])) * y_scale
                page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(0, 0, 0))

            for box in qr_boxes:
                x0 = box["left"] * x_scale
                y0 = box["top"] * y_scale
                x1 = (box["left"] + max(1, box["width"])) * x_scale
                y1 = (box["top"] + max(1, box["height"])) * y_scale
                page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(0, 0, 0))
            page.apply_redactions()

        output = io.BytesIO()
        doc.save(output, garbage=4, deflate=True)
        doc.close()
        flattened_findings = [item for chunk in findings_bucket for item in chunk]
        return {
            "sanitized_text": "\n\n".join(sanitized_pages),
            "sanitized_original_bytes": output.getvalue(),
            "sanitized_original_ext": ".pdf",
            "sanitized_original_media_type": "application/pdf",
            "findings": flattened_findings,
            "summary": _collect_counts(findings_bucket),
            "parser": "pdf_ocr",
        }

    if _looks_like_text(file_bytes):
        raw_text = file_bytes.decode("utf-8", errors="ignore")
        sanitized_text, findings = _sanitize_chunk(raw_text, mode)
        sanitized_ext = ext if ext else ".txt"
        return {
            "sanitized_text": sanitized_text,
            "sanitized_original_bytes": sanitized_text.encode("utf-8"),
            "sanitized_original_ext": sanitized_ext,
            "sanitized_original_media_type": "text/plain",
            "findings": findings,
            "summary": _collect_counts([findings]),
            "parser": "generic_text",
        }

    raise ValueError(f"Unsupported file type: {ext or '[no extension]'}")
